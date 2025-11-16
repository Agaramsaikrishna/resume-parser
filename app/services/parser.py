import os
import json
from pathlib import Path
from uuid import uuid4
from typing import Dict, Any

import fitz  # PyMuPDF
from docx import Document
from werkzeug.utils import secure_filename

from app.utils.logger import setup_logger
from app.services.llm_service import extract_structured_json

logger = setup_logger(os.getenv("LOG_LEVEL", "INFO"))

UPLOAD_DIR = Path(os.getenv("UPLOAD_DIR", "./data/uploads"))
METADATA_FILE = Path(os.getenv("METADATA_FILE", "./data/data.json"))

# Ensure directories and metadata file exist
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
METADATA_FILE.parent.mkdir(parents=True, exist_ok=True)
if not METADATA_FILE.exists():
    METADATA_FILE.write_text(json.dumps({}, indent=2), encoding="utf-8")


def extract_text_from_pdf(path: str) -> str:
    """
    Extract text from a PDF file.

    Args:
        path (str): Path to the PDF file.

    Returns:
        str: Extracted text.
    """
    logger.info("Extracting text from PDF: %s", path)
    text_chunks = []

    try:
        with fitz.open(path) as doc:
            for page in doc:
                text = page.get_text()
                if text:
                    text_chunks.append(text)
    except Exception as exc:
        logger.exception("Failed to read PDF: %s", path)
        raise ValueError(f"Unable to extract text from PDF: {exc}")

    return "\n".join(text_chunks)


def extract_text_from_docx(path: str) -> str:
    """
    Extract text from a DOCX file.

    Args:
        path (str): Path to the DOCX file.

    Returns:
        str: Extracted text.
    """
    logger.info("Extracting text from DOCX: %s", path)
    texts = []

    try:
        doc = Document(path)
        texts.extend([p.text for p in doc.paragraphs if p.text])

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        texts.append(cell.text)
    except Exception as exc:
        logger.exception("Failed to read DOCX: %s", path)
        raise ValueError(f"Unable to extract text from DOCX: {exc}")

    return "\n".join(texts)


def save_metadata(document_id: str, data: dict) -> None:
    """
    Save parsed resume metadata to JSON file.

    Args:
        document_id (str): Unique ID for the resume.
        data (dict): Parsed resume data.
    """
    try:
        with METADATA_FILE.open("r", encoding="utf-8") as f:
            meta = json.load(f)
    except json.JSONDecodeError:
        logger.warning("Metadata file corrupted, creating new.")
        meta = {}

    meta[document_id] = data

    with METADATA_FILE.open("w", encoding="utf-8") as f:
        json.dump(meta, f, indent=2)
    logger.info("Saved metadata for document_id: %s", document_id)


def load_metadata(document_id: str) -> dict:
    """
    Load metadata for a given resume ID.

    Args:
        document_id (str): Unique ID for the resume.

    Returns:
        dict: Metadata or None if not found.
    """
    try:
        with METADATA_FILE.open("r", encoding="utf-8") as f:
            meta = json.load(f)
        return meta.get(document_id)
    except Exception as exc:
        logger.exception("Failed to load metadata for document_id: %s", document_id)
        return None


def process_file(file_path: str, filename: str) -> Dict[str, Any]:
    """
    Process a resume file, extract text, parse with LLM, and save metadata.

    Args:
        file_path (str): Path to the uploaded file.
        filename (str): Original file name.

    Returns:
        dict: Parsed resume data.
    """
    ext = filename.lower().split(".")[-1]

    if ext in ("pdf",):
        text = extract_text_from_pdf(file_path)
    elif ext in ("docx", "doc"):
        text = extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file type. Only PDF and DOCX are allowed.")

    if not text.strip():
        raise ValueError("The uploaded file contains no text.")

    logger.info("Calling LLM to parse extracted text...")
    parsed = extract_structured_json(text)

    # Validate parsed output keys
    expected_keys = {"contact", "summary", "experience", "education", "skills", "certifications"}
    for key in expected_keys:
        if key not in parsed:
            parsed[key] = [] if key in {"experience", "education", "skills", "certifications"} else None

    document_id = str(uuid4())
    result = {
        "document_id": document_id,
        "contact": parsed.get("contact"),
        "summary": parsed.get("summary"),
        "experience": parsed.get("experience", []),
        "education": parsed.get("education", []),
        "skills": parsed.get("skills", []),
        "certifications": parsed.get("certifications", []),
        "raw_text": text[:5000],  # truncate to avoid large storage
    }

    save_metadata(document_id, result)
    logger.info("Completed processing resume: %s", document_id)
    return result


def save_uploaded_file(contents: bytes, filename: str) -> str:
    """
    Safely save uploaded file using a secure filename.

    Args:
        contents (bytes): File content.
        filename (str): Original file name.

    Returns:
        str: Path where file was saved.
    """
    safe_name = secure_filename(filename)
    save_path = UPLOAD_DIR / f"{uuid4()}_{safe_name}"
    try:
        with save_path.open("wb") as f:
            f.write(contents)
        logger.info("Uploaded file saved to: %s", save_path)
        return str(save_path)
    except Exception as exc:
        logger.exception("Failed to save uploaded file: %s", filename)
        raise ValueError(f"Failed to save file: {exc}")
